"""Ingest: get plain text out of a protocol source.

Accepts a `.docx` path (heavy libs lazy-imported) or a pasted `.txt` / `.md`
string. Language-agnostic — it never translates; the demo protocol is Polish and
must come through byte-for-byte in the text.
"""

from __future__ import annotations

import os
import re
import zipfile


def _docx_to_text(path: str) -> str:
    """Extract paragraph text from a .docx.

    Tries python-docx first (handles tables, ordering well); falls back to
    unzipping word/document.xml directly so we have zero hard dependency.
    """
    try:  # lazy import — only pay for it when a .docx is actually ingested
        import docx  # type: ignore

        document = docx.Document(path)
        lines = [p.text for p in document.paragraphs]
        # python-docx keeps table cells out of .paragraphs; pull them in too.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        lines.append(txt)
        return "\n".join(l for l in lines if l is not None)
    except ImportError:
        return _docx_to_text_raw(path)


def _docx_to_text_raw(path: str) -> str:
    """Dependency-free fallback: parse word/document.xml paragraphs by regex."""
    import html

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")

    lines: list[str] = []
    for para in re.split(r"</w:p>", xml):
        # tabs and line breaks inside a run become spaces
        para = re.sub(r"<w:(tab|br)\b[^>]*/?>", " ", para)
        texts = re.findall(r"<w:t\b[^>]*>(.*?)</w:t>", para, flags=re.S)
        line = html.unescape("".join(texts)).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def ingest(source: str) -> str:
    """Return plain text for a protocol.

    `source` may be:
      - a path to a .docx file
      - a path to a .txt / .md file
      - a raw pasted string (treated as the protocol text itself)
    """
    looks_like_path = "\n" not in source and len(source) < 4096
    if looks_like_path and os.path.exists(source):
        lower = source.lower()
        if lower.endswith(".docx"):
            return _docx_to_text(source).strip()
        if lower.endswith((".txt", ".md")):
            with open(source, "r", encoding="utf-8") as fh:
                return fh.read().strip()
        # unknown extension but a real file: read as text
        with open(source, "r", encoding="utf-8") as fh:
            return fh.read().strip()

    # not a path — it's the pasted protocol text
    return source.strip()
